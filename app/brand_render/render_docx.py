#!/usr/bin/env python3
"""
DOCX Native Renderer: IR JSON → .docx per KIWIAI-DECK-GENERATION-STANDARD-V1 §2.5
Quality targets: headings/paragraphs/tables editable, diagrams as high-res PNG images.
Consumes the same IR (deck_ir.json with 10 extension fields) as the pptx renderer.
"""
import json, sys, os, io, tempfile, re

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Constants ──────────────────────────────────────────────────────────
FONT_HEADING = 'Microsoft YaHei'
FONT_BODY    = 'Microsoft YaHei'
# A4 portrait with comfortable margins
PAGE_W = 8.27   # inches (A4 width)
PAGE_H = 11.69  # inches (A4 height)
MARGIN = 0.8    # inches

def hex_to_rgb(h):
    h = h.lstrip('#')
    if len(h) == 3:
        h = h[0]*2 + h[1]*2 + h[2]*2
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# SVG helpers (svg_to_png + diagram generators) live in the shared module so both
# the pptx and docx brand renderers use one environment-adaptive render path.
try:
    from .svg_render import svg_to_png, svg_layers, svg_flow
except ImportError:  # direct CLI run (python render_docx.py ...)
    import sys as _sys
    _sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    from svg_render import svg_to_png, svg_layers, svg_flow
# ═══════════════════════════════════════════════════════════════════════
# DOCX HELPERS
# ═══════════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    """Set cell background color with correct Word semantics.
    Must include w:val="clear" + w:color="auto" + w:fill;
    bare w:fill-only shading causes Word to render solid black background
    (Word interprets missing w:val as solid pattern with foreground w:color=auto=black).
    Documented at: AIAGENCY-ENGINEERING-LESSONS F3-a, KIWIAI-DECK-GENERATION-STANDARD-V1 §5.2."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex.lstrip("#")}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_paragraph(doc, text, style=None, font_size=11, bold=False, color=None,
                         font_name=None, alignment=None, space_after=6, space_before=0):
    """Add a paragraph with styling."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = font_name or FONT_BODY
    if color:
        run.font.color.rgb = hex_to_rgb(color)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    return p


def add_bullet_list(doc, items, font_size=11, color='#333'):
    """Add bulleted list items."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(font_size)
        run.font.name = FONT_BODY
        run.font.color.rgb = hex_to_rgb(color)
    return doc


def add_page_break(doc):
    """Add a page break to start a new section."""
    p = doc.add_paragraph()
    run = p.add_run()
    run._r.append(parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>'))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


# ═══════════════════════════════════════════════════════════════════════
# MAIN RENDERER
# ═══════════════════════════════════════════════════════════════════════

def _render_to_doc(ir, brand_dir=None, work_dir=None):
    """Core renderer: IR dict -> docx Document object.

    Args:
        ir: IR dict (deck_ir.json structure: brand_tokens / cover_reference / slides).
        brand_dir: Directory containing the customer's deck-brand assets (e.g. logo.png).
                   Combined with IR `cover_reference.logo_path` to resolve the logo.
                   No customer slug/name literal lives in this source (P2-3) — the caller
                   supplies the customer directory.
        work_dir: Working dir for temp PNG files (default: tempdir/kiwiai-brand-render).
    """
    work_dir = work_dir or os.path.join(tempfile.gettempdir(), 'kiwiai-brand-render')
    os.makedirs(work_dir, exist_ok=True)
    
    brand = ir['brand_tokens']
    primary = brand['palette']['primary']
    primary_dark = brand['palette'].get('primary_dark', '#005A9B')
    primary_light = brand['palette']['primary_light']
    primary_pale = brand['palette']['primary_pale']
    accent = brand['palette']['accent']
    
    cover_ref = ir['cover_reference']
    company_cn = cover_ref['company_name_cn']
    company_en = cover_ref['company_name_en']
    
    # Short company name (strip prefix city + suffix legal form)
    short_cn = re.sub(r'^(北京|上海|深圳|广州|杭州|成都|武汉|南京|重庆|天津|苏州|西安|长沙|郑州|青岛|大连|厦门|宁波)', '', company_cn)
    short_cn = re.sub(r'(科技|技术|网络|信息|数据|软件|智能)?(有限公司|股份公司|集团公司|有限责任公司|股份有限公司|集团有限公司)$', '', short_cn)
    company_short = short_cn.strip() or company_cn
    
    # Resolve logo path from IR (NO customer slug literal — P2-3 / redline: zero
    # hardcoding). Resolution: brand_dir + IR logo_path > unresolved (rendered w/o logo).
    logo_rel = cover_ref.get('logo_path', '')
    logo_path = None
    if brand_dir and logo_rel:
        logo_path = os.path.join(brand_dir, logo_rel)
    
    # ── Create Document ──────────────────────────────────────────────
    doc = Document()
    
    # Page setup: A4 portrait
    section = doc.sections[0]
    section.page_width = Inches(PAGE_W)
    section.page_height = Inches(PAGE_H)
    section.top_margin = Inches(MARGIN)
    section.bottom_margin = Inches(MARGIN)
    section.left_margin = Inches(MARGIN)
    section.right_margin = Inches(MARGIN)
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(11)
    style.font.color.rgb = hex_to_rgb('#333333')
    
    temp_files = []
    
    # ── Render Each Slide ────────────────────────────────────────────
    slides = ir['slides']
    for idx, slide_data in enumerate(slides):
        if idx > 0:
            add_page_break(doc)
        
        stype = slide_data['type']
        
        # ═══════════════════ COVER ═══════════════════
        if stype == 'cover':
            # Logo centered at top
            if logo_path and os.path.exists(logo_path):
                p_logo = doc.add_paragraph()
                p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_logo = p_logo.add_run()
                run_logo.add_picture(logo_path, width=Inches(1.5), height=Inches(0.9))
                p_logo.paragraph_format.space_after = Pt(12)
            
            # Company name CN
            add_styled_paragraph(doc, company_cn, font_size=15, bold=True,
                                 color='#1A1A2E', font_name=FONT_HEADING,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
            # Company name EN
            add_styled_paragraph(doc, company_en, font_size=11,
                                 color='#666666', font_name=FONT_BODY,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
            
            # Divider line (thin rectangle via paragraph border is simpler)
            p_div = doc.add_paragraph()
            p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_div.paragraph_format.space_after = Pt(16)
            # Add a bottom border to simulate divider
            pPr = p_div._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="{primary.lstrip("#")}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            
            # Title
            title = slide_data.get('title', '')
            if title:
                add_styled_paragraph(doc, title, font_size=28, bold=True,
                                     color=primary, font_name=FONT_HEADING,
                                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
            
            # Subtitle
            subtitle = slide_data.get('subtitle', '')
            if subtitle:
                add_styled_paragraph(doc, subtitle, font_size=13,
                                     color='#666666', font_name=FONT_BODY,
                                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
            
            # Footer
            footer_text = slide_data.get('footer', company_short + ' · 2026')
            add_styled_paragraph(doc, footer_text, font_size=10,
                                 color='#999999', font_name=FONT_BODY,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
        
        # ═══════════════════ END ═══════════════════
        elif stype == 'end':
            # Logo centered
            if logo_path and os.path.exists(logo_path):
                p_logo = doc.add_paragraph()
                p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_logo = p_logo.add_run()
                run_logo.add_picture(logo_path, width=Inches(1.2), height=Inches(0.72))
                p_logo.paragraph_format.space_after = Pt(12)
            
            # Company name CN (brand color)
            add_styled_paragraph(doc, company_cn, font_size=15, bold=True,
                                 color=primary, font_name=FONT_HEADING,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
            # Company name EN
            add_styled_paragraph(doc, company_en, font_size=11,
                                 color='#666666', font_name=FONT_BODY,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
            
            # Divider
            p_div = doc.add_paragraph()
            p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_div.paragraph_format.space_after = Pt(14)
            pPr = p_div._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="{primary.lstrip("#")}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            
            # 感谢聆听
            title = slide_data.get('title', '感谢聆听')
            add_styled_paragraph(doc, title, font_size=28, bold=True,
                                 color=primary, font_name=FONT_HEADING,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
            
            # Subtitle (方案名)
            subtitle = slide_data.get('subtitle', '')
            if subtitle:
                add_styled_paragraph(doc, subtitle, font_size=14,
                                     color='#666666', font_name=FONT_HEADING,
                                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
            
            # Contact info
            contact = slide_data.get('contact', '')
            if contact:
                parts = contact.split('|')
                for part in parts:
                    part = part.strip()
                    if part:
                        add_styled_paragraph(doc, part, font_size=11,
                                             color='#999999', font_name=FONT_BODY,
                                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        
        # ═══════════════════ CONTENT ═══════════════════
        elif stype == 'content':
            tag = slide_data.get('tag', '')
            title = slide_data.get('title', '')
            
            # Section tag as small colored label
            if tag:
                add_styled_paragraph(doc, tag, font_size=9, bold=True,
                                     color=primary, font_name=FONT_HEADING,
                                     space_after=2)
            
            # Title as Heading 1
            if title:
                add_styled_paragraph(doc, title, font_size=20, bold=True,
                                     color='#1A1A2E', font_name=FONT_HEADING,
                                     space_after=14)
                # Add a thin colored line under heading
                p_line = doc.add_paragraph()
                p_line.paragraph_format.space_after = Pt(10)
                p_line.paragraph_format.space_before = Pt(0)
                pPr = p_line._p.get_or_add_pPr()
                pBdr = parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="{primary.lstrip("#")}"/>'
                    f'</w:pBdr>'
                )
                pPr.append(pBdr)
            
            # ── Agenda items ──
            if 'agenda_items' in slide_data:
                items = slide_data['agenda_items']
                for i, item in enumerate(items):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(6)
                    # Number badge
                    run_num = p.add_run(f'  {str(i+1).zfill(2)}  ')
                    run_num.font.size = Pt(11)
                    run_num.font.bold = True
                    run_num.font.color.rgb = hex_to_rgb('#FFFFFF')
                    # Underline-style number prefix
                    pPr = p._p.get_or_add_pPr()
                    
                    run_name = p.add_run(f'  {item}')
                    run_name.font.size = Pt(13)
                    run_name.font.bold = True
                    run_name.font.name = FONT_BODY
                    run_name.font.color.rgb = hex_to_rgb('#1A1A2E')
            
            # ── Bullets (+ optional highlight box) ──
            elif 'bullets' in slide_data:
                bullets = slide_data['bullets']
                has_two_col = slide_data.get('layout') == 'two-column'
                highlight = slide_data.get('highlight', '')
                stats = slide_data.get('stats', [])
                
                if has_two_col and (highlight or stats):
                    # Two-column: bullets on left, highlight box on right via table
                    tbl = doc.add_table(rows=1, cols=2)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    # Left: bullets
                    left_cell = tbl.cell(0, 0)
                    for bullet in bullets:
                        p = left_cell.add_paragraph()
                        p.paragraph_format.space_after = Pt(6)
                        run = p.add_run(f'● {bullet}')
                        run.font.size = Pt(12)
                        run.font.name = FONT_BODY
                        run.font.color.rgb = hex_to_rgb('#333333')
                    # Right: highlight box
                    right_cell = tbl.cell(0, 1)
                    set_cell_shading(right_cell, primary_pale.lstrip('#'))
                    # Accent left border via shading only; add company promise
                    p_hdr = right_cell.add_paragraph()
                    run_hdr = p_hdr.add_run(f'{company_short}承诺')
                    run_hdr.font.size = Pt(13)
                    run_hdr.font.bold = True
                    run_hdr.font.name = FONT_HEADING
                    run_hdr.font.color.rgb = hex_to_rgb(primary)
                    p_hdr.paragraph_format.space_after = Pt(8)
                    
                    # Strip stats values from highlight text for clean display
                    hl_clean = highlight
                    for s in stats:
                        hl_clean = hl_clean.replace(' ' + s['value'] + ' ' + s['label'], '')
                    hl_clean = hl_clean.replace('  ', ' ').strip()
                    
                    p_body = right_cell.add_paragraph()
                    run_body = p_body.add_run(hl_clean)
                    run_body.font.size = Pt(11)
                    run_body.font.name = FONT_BODY
                    run_body.font.color.rgb = hex_to_rgb('#333333')
                    p_body.paragraph_format.space_after = Pt(10)
                    
                    # Stats
                    for s in stats:
                        p_stat = right_cell.add_paragraph()
                        run_val = p_stat.add_run(s['value'] + '  ')
                        run_val.font.size = Pt(18)
                        run_val.font.bold = True
                        run_val.font.color.rgb = hex_to_rgb(primary)
                        run_val.font.name = FONT_HEADING
                        run_lbl = p_stat.add_run(s['label'])
                        run_lbl.font.size = Pt(10)
                        run_lbl.font.color.rgb = hex_to_rgb('#666666')
                        run_lbl.font.name = FONT_BODY
                        p_stat.paragraph_format.space_after = Pt(2)
                    
                    # Remove first empty paragraph that add_table creates
                    for cell in [left_cell, right_cell]:
                        if cell.paragraphs and cell.paragraphs[0].text == '':
                            p_elem = cell.paragraphs[0]._p
                            p_elem.getparent().remove(p_elem)
                else:
                    add_bullet_list(doc, bullets, font_size=12)
            
            # ── Cards ──
            elif 'cards' in slide_data:
                cards = slide_data['cards']
                n = len(cards)
                if n <= 4:
                    # Use a table: 1 row, n columns for card layout
                    tbl = doc.add_table(rows=1, cols=n)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for ci, card in enumerate(cards):
                        cell = tbl.cell(0, ci)
                        icon = card.get('icon', '')
                        card_title = card.get('title', '')
                        body = card.get('body', '')
                        
                        if icon:
                            p_icon = cell.add_paragraph()
                            run_icon = p_icon.add_run(icon)
                            run_icon.font.size = Pt(22)
                            p_icon.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p_icon.paragraph_format.space_after = Pt(4)
                        
                        p_ttl = cell.add_paragraph()
                        run_ttl = p_ttl.add_run(card_title)
                        run_ttl.font.size = Pt(13)
                        run_ttl.font.bold = True
                        run_ttl.font.name = FONT_HEADING
                        run_ttl.font.color.rgb = hex_to_rgb('#1A1A2E')
                        p_ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_ttl.paragraph_format.space_after = Pt(4)
                        
                        if body:
                            p_bd = cell.add_paragraph()
                            run_bd = p_bd.add_run(body)
                            run_bd.font.size = Pt(10)
                            run_bd.font.name = FONT_BODY
                            run_bd.font.color.rgb = hex_to_rgb('#666666')
                            p_bd.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Remove first empty paragraph
                        if cell.paragraphs and cell.paragraphs[0].text == '' and len(cell.paragraphs) > 1:
                            p_elem = cell.paragraphs[0]._p
                            p_elem.getparent().remove(p_elem)
                else:
                    # Fallback: list cards vertically
                    for card in cards:
                        icon = card.get('icon', '')
                        card_title = card.get('title', '')
                        body = card.get('body', '')
                        p = doc.add_paragraph()
                        if icon:
                            run_icon = p.add_run(icon + ' ')
                            run_icon.font.size = Pt(16)
                        if card_title:
                            run_ttl = p.add_run(card_title)
                            run_ttl.font.size = Pt(13)
                            run_ttl.font.bold = True
                            run_ttl.font.name = FONT_HEADING
                            run_ttl.font.color.rgb = hex_to_rgb('#1A1A2E')
                        if body:
                            p2 = doc.add_paragraph()
                            run_bd = p2.add_run(body)
                            run_bd.font.size = Pt(10)
                            run_bd.font.name = FONT_BODY
                            run_bd.font.color.rgb = hex_to_rgb('#666666')
            
            # ── Timeline ──
            elif 'timeline' in slide_data:
                timeline = slide_data['timeline']
                n = len(timeline)
                # Table: header row with phases, content row with title+description
                tbl = doc.add_table(rows=1, cols=n)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                for ti, item in enumerate(timeline):
                    cell = tbl.cell(0, ti)
                    phase = item.get('phase', '')
                    item_title = item.get('title', '')
                    desc = item.get('description', '')
                    
                    # Phase label
                    p_phase = cell.add_paragraph()
                    run_ph = p_phase.add_run(f'Phase {phase}')
                    run_ph.font.size = Pt(9)
                    run_ph.font.bold = True
                    run_ph.font.color.rgb = hex_to_rgb(primary)
                    run_ph.font.name = FONT_HEADING
                    p_phase.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_phase.paragraph_format.space_after = Pt(4)
                    
                    # Title
                    p_ttl = cell.add_paragraph()
                    run_t = p_ttl.add_run(item_title)
                    run_t.font.size = Pt(12)
                    run_t.font.bold = True
                    run_t.font.name = FONT_HEADING
                    run_t.font.color.rgb = hex_to_rgb('#1A1A2E')
                    p_ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_ttl.paragraph_format.space_after = Pt(4)
                    
                    # Description
                    p_desc = cell.add_paragraph()
                    run_d = p_desc.add_run(desc)
                    run_d.font.size = Pt(9)
                    run_d.font.name = FONT_BODY
                    run_d.font.color.rgb = hex_to_rgb('#666666')
                    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Set cell shading
                    set_cell_shading(cell, primary_pale.lstrip('#'))
                    
                    if cell.paragraphs and cell.paragraphs[0].text == '' and len(cell.paragraphs) > 1:
                        p_elem = cell.paragraphs[0]._p
                        p_elem.getparent().remove(p_elem)
        
        # ═══════════════════ DIAGRAM ═══════════════════
        elif stype == 'diagram':
            tag = slide_data.get('tag', '')
            title = slide_data.get('title', '')
            
            if tag:
                add_styled_paragraph(doc, tag, font_size=9, bold=True,
                                     color=primary, font_name=FONT_HEADING, space_after=2)
            if title:
                add_styled_paragraph(doc, title, font_size=20, bold=True,
                                     color='#1A1A2E', font_name=FONT_HEADING, space_after=8)
            
            # Generate SVG from IR data
            svg_kind = slide_data.get('svg_kind', '')
            svg_content = None
            
            if svg_kind == 'layers':
                layers = slide_data.get('layers', [])
                if layers:
                    svg_content = svg_layers(title, layers, primary, primary_pale)
            elif svg_kind == 'flow':
                steps = slide_data.get('steps', [])
                if steps:
                    svg_content = svg_flow(title, steps, primary, primary_light)
            
            if svg_content:
                png_path = os.path.join(work_dir, f'_diagram_{idx}.png')
                temp_files.append(png_path)
                svg_to_png(svg_content, png_path, width=1280, height=600)
                # Embed PNG
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                run_img.add_picture(png_path, width=Inches(PAGE_W - 2 * MARGIN))
                p_img.paragraph_format.space_after = Pt(6)
            else:
                # No SVG data — this should not happen if IR is well-formed
                add_styled_paragraph(doc, f'[图示: {title}] (缺少 SVG 数据)', font_size=11,
                                     color='#CC0000', alignment=WD_ALIGN_PARAGRAPH.CENTER)
            
            # Insight text
            insight = slide_data.get('insight', '')
            if insight:
                p_ins = doc.add_paragraph()
                p_ins.paragraph_format.space_before = Pt(8)
                # Split into sentences
                sentences = re.split(r'(?<=[。])\s*', insight)
                for sent in sentences:
                    sent = sent.strip()
                    if sent:
                        run_s = p_ins.add_run(sent + '  ')
                        run_s.font.size = Pt(9)
                        run_s.font.name = FONT_BODY
                        run_s.font.color.rgb = hex_to_rgb('#888888')
        
        # ═══════════════════ CHART ═══════════════════
        elif stype == 'chart':
            tag = slide_data.get('tag', '')
            title = slide_data.get('title', '')
            
            if tag:
                add_styled_paragraph(doc, tag, font_size=9, bold=True,
                                     color=primary, font_name=FONT_HEADING, space_after=2)
            if title:
                add_styled_paragraph(doc, title, font_size=20, bold=True,
                                     color='#1A1A2E', font_name=FONT_HEADING, space_after=10)
            
            # Chart data as editable Word table
            labels = slide_data.get('labels', [])
            values = slide_data.get('values', [])
            unit = slide_data.get('unit', '%')
            
            if labels and values:
                n = len(labels)
                tbl = doc.add_table(rows=n + 1, cols=2)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                tbl.style = 'Table Grid'
                
                # Header row
                hdr_label = tbl.cell(0, 0)
                hdr_val = tbl.cell(0, 1)
                set_cell_shading(hdr_label, primary.lstrip('#'))
                set_cell_shading(hdr_val, primary.lstrip('#'))
                
                p_hl = hdr_label.add_paragraph()
                run_hl = p_hl.add_run('指标')
                run_hl.font.size = Pt(11)
                run_hl.font.bold = True
                run_hl.font.name = FONT_HEADING
                run_hl.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                p_hl.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                p_hv = hdr_val.add_paragraph()
                run_hv = p_hv.add_run(f'数值 ({unit})')
                run_hv.font.size = Pt(11)
                run_hv.font.bold = True
                run_hv.font.name = FONT_HEADING
                run_hv.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                p_hv.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Remove default empty paragraphs
                for cell in [hdr_label, hdr_val]:
                    if cell.paragraphs and len(cell.paragraphs) > 1 and cell.paragraphs[0].text == '':
                        p_elem = cell.paragraphs[0]._p
                        p_elem.getparent().remove(p_elem)
                
                # Data rows
                for i, (label, value) in enumerate(zip(labels, values)):
                    row_idx = i + 1
                    cell_l = tbl.cell(row_idx, 0)
                    cell_v = tbl.cell(row_idx, 1)
                    
                    # Alternate row shading
                    if i % 2 == 0:
                        set_cell_shading(cell_l, 'F8FAFB')
                        set_cell_shading(cell_v, 'F8FAFB')
                    
                    p_l = cell_l.add_paragraph()
                    run_l = p_l.add_run(label)
                    run_l.font.size = Pt(11)
                    run_l.font.name = FONT_BODY
                    run_l.font.color.rgb = hex_to_rgb('#333333')
                    
                    p_v = cell_v.add_paragraph()
                    run_v = p_v.add_run(f'{value}{unit}')
                    run_v.font.size = Pt(11)
                    run_v.font.bold = True
                    run_v.font.name = FONT_HEADING
                    run_v.font.color.rgb = hex_to_rgb(primary)
                    p_v.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Remove default empty paras
                    for cell in [cell_l, cell_v]:
                        if cell.paragraphs and len(cell.paragraphs) > 1 and cell.paragraphs[0].text == '':
                            p_elem = cell.paragraphs[0]._p
                            p_elem.getparent().remove(p_elem)
            else:
                add_styled_paragraph(doc, '[图表: 无数据]', font_size=11,
                                     color='#999999', alignment=WD_ALIGN_PARAGRAPH.CENTER)
            
            # Insight text
            insight = slide_data.get('insight', '')
            if insight:
                p_ins = doc.add_paragraph()
                p_ins.paragraph_format.space_before = Pt(10)
                sentences = re.split(r'(?<=[。])\s*', insight)
                for sent in sentences:
                    sent = sent.strip()
                    if sent:
                        run_s = p_ins.add_run(sent + '  ')
                        run_s.font.size = Pt(9)
                        run_s.font.name = FONT_BODY
                        run_s.font.color.rgb = hex_to_rgb('#888888')
    
    # ── Save ──────────────────────────────────────────────────────────
    # Clean up temp files
    for tf in temp_files:
        try:
            if os.path.exists(tf):
                os.unlink(tf)
        except Exception:
            pass
    
    return doc


def render_docx_from_ir(ir, brand_dir=None, work_dir=None):
    """Library entry: IR dict -> docx bytes (used by artifact_generator server)."""
    doc = _render_to_doc(ir, brand_dir=brand_dir, work_dir=work_dir)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_docx(ir_path, output_path=None, brand_dir=None, work_dir=None):
    """CLI entry: IR file -> output .docx file."""
    with open(ir_path, 'r', encoding='utf-8') as f:
        ir = json.load(f)
    doc = _render_to_doc(ir, brand_dir=brand_dir, work_dir=work_dir)
    output_path = output_path or 'output.docx'
    doc.save(output_path)
    return doc


# ═══════════════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python3 render_docx.py <deck_ir.json> [output.docx] [brand_dir]")
        sys.exit(1)
    
    ir_path = sys.argv[1]
    output = sys.argv[2] if len(sys.argv) > 2 else 'output.docx'
    brand_dir = sys.argv[3] if len(sys.argv) > 3 else None
    
    render_docx(ir_path, output, brand_dir)
    print(f'DOCX saved: {output}')
